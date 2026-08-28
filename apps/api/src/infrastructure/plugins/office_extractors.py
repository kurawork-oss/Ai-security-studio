"""Office / PDF extractor plugins (optional deps: pypdf, python-docx, openpyxl).

Each is registered as an available extractor only when its library is
importable; otherwise a stub manifest (available=False) is advertised instead.
Extracted text always flows through the PII engine before leaving the system.
"""

from __future__ import annotations

import io

from ...core.errors import ValidationError
from ...domain.plugins import PluginManifest

_DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_CT = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


class DocxExtractor:
    manifest = PluginManifest("docx", "extractor", description="Word (.docx) to text",
                              content_types=(_DOCX_CT,))

    def supports(self, content_type: str) -> bool:
        return content_type in self.manifest.content_types

    def extract(self, content: bytes, content_type: str) -> str:
        import docx  # type: ignore

        try:
            doc = docx.Document(io.BytesIO(content))
        except Exception as exc:
            raise ValidationError("Could not parse .docx content") from exc
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)


class XlsxExtractor:
    manifest = PluginManifest("xlsx", "extractor", description="Excel (.xlsx) to text",
                              content_types=(_XLSX_CT,))

    def supports(self, content_type: str) -> bool:
        return content_type in self.manifest.content_types

    def extract(self, content: bytes, content_type: str) -> str:
        import openpyxl  # type: ignore

        try:
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        except Exception as exc:
            raise ValidationError("Could not parse .xlsx content") from exc
        out: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    out.append(" ".join(cells))
        return "\n".join(out)


class PdfExtractor:
    manifest = PluginManifest("pdf", "extractor", description="PDF (text layer) to text",
                              content_types=("application/pdf",))

    def supports(self, content_type: str) -> bool:
        return content_type in self.manifest.content_types

    def extract(self, content: bytes, content_type: str) -> str:
        from pypdf import PdfReader  # type: ignore
        from pypdf.errors import PdfError  # type: ignore

        try:
            reader = PdfReader(io.BytesIO(content))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        except (PdfError, OSError, ValueError) as exc:
            raise ValidationError("Could not parse PDF content") from exc


_CANDIDATES = [(DocxExtractor, "docx"), (XlsxExtractor, "openpyxl"), (PdfExtractor, "pypdf")]


def available_office_extractors() -> list:
    out = []
    for cls, module in _CANDIDATES:
        try:
            __import__(module)
            out.append(cls())
        except ImportError:
            pass
    return out


def missing_office_manifests() -> list[PluginManifest]:
    out = []
    for cls, module in _CANDIDATES:
        try:
            __import__(module)
        except ImportError:
            m = cls.manifest
            out.append(
                PluginManifest(m.key, m.category, description=m.description,
                               content_types=m.content_types, available=False)
            )
    return out
